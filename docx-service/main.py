from fastapi import FastAPI
from pydantic import BaseModel
from docx import Document
from fastapi.responses import Response
import io

app = FastAPI()

class SOWRequest(BaseModel):
    title: str
    unit: str
    content: dict

@app.post("/generate")
async def generate_docx(sow: SOWRequest):
    doc = Document()
    doc.add_heading(sow.title, 0)
    doc.add_paragraph(f"Unit: {sow.unit}")
    
    for section, text in sow.content.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(text)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    headers = {
        'Content-Disposition': f'attachment; filename="sow-{sow.title}.docx"'
    }

    return Response(content=file_stream.getvalue(), media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers=headers)